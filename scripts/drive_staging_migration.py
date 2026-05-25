#!/usr/bin/env python3
"""Load HODOM Drive sources into PostgreSQL staging tables.

This script intentionally writes patient-identifiable extracted rows only to the
local database. Generated SQL and downloaded raw files stay under .tmp/.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import subprocess
import sys
import unicodedata
from datetime import date, datetime, time
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook

if str(ROOT := Path(__file__).resolve().parents[1]) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.drive_consolidation import MANIFEST as METRICS_MANIFEST


OUT_DIR = ROOT / ".tmp" / "drive-migration"
RAW_DIR = OUT_DIR / "raw"
LOAD_SQL = OUT_DIR / "drive-staging-load.sql"
HANDOVER_MANIFEST = OUT_DIR / "handover-manifest.json"
DB_URL = "postgresql://hodom:hodom@localhost:5555/hodom"
PHASE = "drive_staging_migration_2026_05_25"

METRICS_FOLDER_URL = "https://drive.google.com/drive/folders/1P6ur0dHwmECADUluzwfVvA5g9uTO9oVe?usp=sharing"
ROUTES_FOLDER_URL = "https://drive.google.com/drive/folders/1Jv5ZbSKBMeHszSzbdn0__dXpOr-wKznl?usp=sharing"
HANDOVER_FOLDER_URL = "https://drive.google.com/drive/folders/13Z2fsVSrjl9IiZ4vIvMwltK_MMEMLTGC?usp=sharing"

METRICS_SPEC = ROOT / "docs" / "specs" / "metricas-hodom" / "metricas-anuales-atenciones-usuarios-2026-05-25.json"
CONSOLIDATION_SPEC = ROOT / "docs" / "specs" / "metricas-hodom" / "consolidacion-fuentes-drive-2026-05-25.json"

MONTHS = {
    "ENERO": 1,
    "FEBRERO": 2,
    "MARZO": 3,
    "ABRIL": 4,
    "MAYO": 5,
    "JUNIO": 6,
    "JULIO": 7,
    "AGOSTO": 8,
    "SEPTIEMBRE": 9,
    "SETIEMBRE": 9,
    "SEPT": 9,
    "OCTUBRE": 10,
    "OCT": 10,
    "NOVIEMBRE": 11,
    "DICIEMBRE": 12,
}

ROUTE_SOURCES: list[dict[str, Any]] = [
    {"drive_id": "1IUpd27NDjFJ8YIw24uiwfEJfXyuT4CLp", "path": "rutas 2024/VISITAS ENERO.xlsx", "title": "VISITAS ENERO.xlsx", "year": 2024, "month": 1},
    {"drive_id": "1il8Txeg2z9RglNmtm_5uJzp7GZP6yXz_", "path": "rutas 2024/VISITAS FEB.xlsx", "title": "VISITAS FEB.xlsx", "year": 2024, "month": 2},
    {"drive_id": "1O2472ZxWbvHOQXZXxDv_gYKJ3GGeFJXO", "path": "rutas 2024/MARZO VISITAS.xlsx", "title": "MARZO VISITAS.xlsx", "year": 2024, "month": 3},
    {"drive_id": "1bnwMzRhMrf3kkIGhKM1y-VMDr7_TuZL3", "path": "rutas 2024/VISITAS ABRIL.xlsx", "title": "VISITAS ABRIL.xlsx", "year": 2024, "month": 4},
    {"drive_id": "1oQqpww9WJX73q_7k7sIWCb1-zPvHx5mV", "path": "rutas 2024/VISITAS MAYO.xlsx", "title": "VISITAS MAYO.xlsx", "year": 2024, "month": 5},
    {"drive_id": "1zLBpS6e1o215gUenlv8o0bpbfr3KBm0S", "path": "rutas 2024/VISITAS JUNIO.xlsx", "title": "VISITAS JUNIO.xlsx", "year": 2024, "month": 6},
    {"drive_id": "1JsoySwQvbL0JOypU3FbKERMRSCjdTzbA", "path": "rutas 2024/VISITAS JULIO.xlsx", "title": "VISITAS JULIO.xlsx", "year": 2024, "month": 7},
    {"drive_id": "15rIH_dvzGG8SqM1FshOQJc3C_R1PMd7Q", "path": "rutas 2024/VISITAS AGOSTO.xlsx", "title": "VISITAS AGOSTO.xlsx", "year": 2024, "month": 8},
    {"drive_id": "1C8aEcgmTrOQ1EnJfZrraiEQ3E8zQKxFg", "path": "rutas 2024/VISITAS SEPTIEMBRE.xlsx", "title": "VISITAS SEPTIEMBRE.xlsx", "year": 2024, "month": 9},
    {"drive_id": "1JelWNFySCwlAWiiIdtZlrGk_zqnnFu10", "path": "rutas 2024/VISITAS OCTUBRE 24.xlsx", "title": "VISITAS OCTUBRE 24.xlsx", "year": 2024, "month": 10},
    {"drive_id": "1_qvQrhMhpqtNDOK7IHfGWoj1VITq_kcX", "path": "rutas 2024/VISITAS NOVIEMBRE 24.xlsx", "title": "VISITAS NOVIEMBRE 24.xlsx", "year": 2024, "month": 11},
    {"drive_id": "1u4x-_I1G1a-_mLDRpE4aFiy-Lybo4L_a", "path": "rutas 2024/VISITAS DICIEMBRE 2024.xlsx", "title": "VISITAS DICIEMBRE 2024.xlsx", "year": 2024, "month": 12},
    {"drive_id": "1PL77Odq_lfRA2asAm9npNbTFuMOE0_Vu", "path": "RUTAS 2025/VISITAS ENERO 25.xlsx", "title": "VISITAS ENERO 25.xlsx", "year": 2025, "month": 1},
    {"drive_id": "1fVaMobQD2kqBcdqCQ55xO_iGUDGYW4cU", "path": "RUTAS 2025/VISITAS FEBRERO 25.xlsx", "title": "VISITAS FEBRERO 25.xlsx", "year": 2025, "month": 2},
    {"drive_id": "14ZylrS1FeMsEfVTWAJPRSResnnPRAsis", "path": "RUTAS 2025/VISITAS MARZO 25.xlsx", "title": "VISITAS MARZO 25.xlsx", "year": 2025, "month": 3},
    {"drive_id": "1GTtWRztpMe1ALDsIq9cUobj5aVqnXD4t", "path": "RUTAS 2025/VISITAS ABRIL 25.xlsx", "title": "VISITAS ABRIL 25.xlsx", "year": 2025, "month": 4},
    {"drive_id": "1eQeEBuumRozX5UZwM9usfT-9pXrGKHyw", "path": "RUTAS 2025/VISITAS MAYO", "title": "VISITAS MAYO", "year": 2025, "month": 5},
    {"drive_id": "1ydRZfOfg08bzv_xMNTyHYrFNwj9_8PIn", "path": "RUTAS 2025/VISITAS JUNIO", "title": "VISITAS JUNIO", "year": 2025, "month": 6},
    {"drive_id": "1QG49TK19YYWag0x7BMtEUYk46wCXoKet", "path": "RUTAS 2025/VISITAS JULIO 25.xlsx", "title": "VISITAS JULIO 25.xlsx", "year": 2025, "month": 7},
    {"drive_id": "1OOzVRHowdt8jhHhR9JWR4swWYYFVXbrF", "path": "RUTAS 2025/VISITAS AGOSTO", "title": "VISITAS AGOSTO", "year": 2025, "month": 8},
    {"drive_id": "1DD5aSGn5oXr4x8Zba0AfYvke3aor-AYQ", "path": "RUTAS 2025/VISITAS SEPTIEMBRE", "title": "VISITAS SEPTIEMBRE", "year": 2025, "month": 9},
    {"drive_id": "10lW8WDSY7KIdLp4EDOPB3e1YeGG-I02m", "path": "RUTAS 2025/VISITAS OCTUBRE 25.xlsx", "title": "VISITAS OCTUBRE 25.xlsx", "year": 2025, "month": 10},
    {"drive_id": "1mXmzhyl0SEQOyEuDZU37eK_EAtW3eF3B", "path": "RUTAS 2025/VISITAS NOVIEMBRE 25", "title": "VISITAS NOVIEMBRE 25", "year": 2025, "month": 11},
    {"drive_id": "16Xx2AJAfU3_YFLI2FgP-sXwUy9tUBF6H", "path": "RUTAS 2025/VISITAS DICIEMBRE 25.xlsx", "title": "VISITAS DICIEMBRE 25.xlsx", "year": 2025, "month": 12},
    {"drive_id": "1O4Hahjbg0rbl6jRbMtIRCUlmLmrupiH8", "path": "ENERO 2026.xlsx", "title": "ENERO 2026", "year": 2026, "month": 1},
    {"drive_id": "1NOMPBfDKPSnI7_CGCTvmTpMc-PXb0wFr", "path": "FEBRERO 2026.xlsx", "title": "FEBRERO 2026", "year": 2026, "month": 2},
    {"drive_id": "10lHEj-3_V6ZJb-V0kKI7IecNsqO_oCmT", "path": "MARZO 2026.xlsx", "title": "MARZO 2026", "year": 2026, "month": 3},
    {"drive_id": "1HX5NlbjYQUiLD5ZJmnsHsAghUcJgjaFZ4lqaPBePXjQ", "path": "ABRIL 2026", "title": "ABRIL 2026", "year": 2026, "month": 4, "native_google": True},
    {"drive_id": "1V2LgAUmcNggcMqnuU07Y0en-44nQ5FpxxSkdOemN9Ow", "path": "MAYO 2026", "title": "MAYO 2026", "year": 2026, "month": 5, "native_google": True},
]


def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    text = unicodedata.normalize("NFKD", str(value))
    text = "".join(char for char in text if not unicodedata.combining(char))
    return " ".join(text.upper().split())


def clean_text(value: Any) -> str | None:
    if value is None:
        return None
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    text = str(value).strip()
    return text or None


def month_number_from_title(title: str) -> int | None:
    normalized = normalize_text(title)
    for name, number in MONTHS.items():
        if re.search(rf"\b{name}\b", normalized):
            return number
    return None


def make_id(prefix: str, value: str) -> str:
    return f"{prefix}_{hashlib.sha256(value.encode('utf-8')).hexdigest()[:16]}"


def safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("_")


def sql_literal(value: Any) -> str:
    if value is None:
        return "NULL"
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    if isinstance(value, (int, float)):
        return str(value)
    text = str(value).replace("'", "''")
    return f"'{text}'"


def json_literal(value: Any) -> str:
    return f"{sql_literal(json.dumps(value, ensure_ascii=False, sort_keys=True))}::jsonb"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def source_url(drive_id: str, native_google: bool = False) -> str:
    if native_google:
        return f"https://docs.google.com/spreadsheets/d/{drive_id}/edit"
    return f"https://drive.google.com/file/d/{drive_id}/view"


def local_path_for(source: dict[str, Any]) -> Path:
    name = safe_name(source["path"])
    if not name.lower().endswith((".xlsx", ".docx")):
        name += ".xlsx"
    return RAW_DIR / name


def download_drive_file(source: dict[str, Any], force: bool = False) -> Path:
    RAW_DIR.mkdir(parents=True, exist_ok=True)
    target = local_path_for(source)
    if target.exists() and target.stat().st_size > 0 and not force:
        return target
    if source.get("native_google"):
        url = f"https://docs.google.com/spreadsheets/d/{source['drive_id']}/export?format=xlsx"
    else:
        url = f"https://drive.google.com/uc?export=download&id={source['drive_id']}"
    subprocess.run(["curl", "-L", "-s", "-o", str(target), url], check=True)
    return target


def parse_date_from_sheet_name(sheet_name: str, fallback_year: int | None, fallback_month: int | None) -> date | None:
    if fallback_year is None:
        return None
    normalized = normalize_text(sheet_name)
    match = re.search(r"(\d{1,2})\s*[.\-/ ]\s*(\d{1,2})", normalized)
    if not match:
        match = re.search(r"(\d{2})(\d{2})", normalized)
    if not match:
        return None
    day = int(match.group(1))
    month = int(match.group(2)) if match.group(2) else fallback_month
    if not month:
        return None
    try:
        return date(fallback_year, month, day)
    except ValueError:
        return None


def parse_sheet_date(sheet_name: str, first_row: list[Any], fallback_year: int | None, fallback_month: int | None) -> date | None:
    for value in first_row[:3]:
        if isinstance(value, datetime):
            return value.date()
    return parse_date_from_sheet_name(sheet_name, fallback_year, fallback_month)


def time_text(value: Any) -> str | None:
    if isinstance(value, time):
        return value.strftime("%H:%M")
    if isinstance(value, datetime):
        return value.strftime("%H:%M")
    return clean_text(value)


def row_has_route_header(values: list[Any]) -> bool:
    labels = [normalize_text(value) for value in values[:12]]
    joined = " ".join(labels)
    return "PACIENTE" in joined or ("MEDICO" in joined and ("KINE" in joined or "KINESIOLOGO" in joined))


def is_skippable_route_row(values: list[Any]) -> bool:
    joined = normalize_text(" ".join(str(value) for value in values if value is not None))
    if not joined:
        return True
    if "COLACION" in joined:
        return True
    if joined.startswith("PACIENTES A VISITAR"):
        return True
    return False


def parse_route_workbook(path: Path, source: dict[str, Any]) -> list[dict[str, Any]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    rows: list[dict[str, Any]] = []
    for sheet in workbook.worksheets:
        sheet_rows = [list(row[:12]) for row in sheet.iter_rows(values_only=True)]
        if not sheet_rows:
            continue
        visit_date = parse_sheet_date(sheet.title, sheet_rows[0], source.get("year"), source.get("month"))
        header_index = next((idx for idx, values in enumerate(sheet_rows[:8]) if row_has_route_header(values)), None)
        if header_index is None:
            continue
        for row_index, values in enumerate(sheet_rows[header_index + 1 :], start=header_index + 2):
            padded = values + [None] * (12 - len(values))
            if is_skippable_route_row(padded):
                continue
            patient_name = clean_text(padded[7])
            service_text = clean_text(padded[8])
            address_text = clean_text(padded[9])
            if not patient_name or (not service_text and not address_text):
                continue
            professionals = {
                "medico": clean_text(padded[2]),
                "fono": clean_text(padded[3]),
                "kine": clean_text(padded[4]),
                "enfermera": clean_text(padded[5]),
                "tens": clean_text(padded[6]),
            }
            professionals = {key: value for key, value in professionals.items() if value}
            visit_date_text = visit_date.isoformat() if visit_date else None
            route_visit_id = make_id(
                "drv_visit",
                "|".join(
                    [
                        source["drive_id"],
                        sheet.title,
                        str(row_index),
                        visit_date_text or "",
                        patient_name,
                        service_text or "",
                    ]
                ),
            )
            rows.append(
                {
                    "route_visit_id": route_visit_id,
                    "drive_id": source["drive_id"],
                    "source_path": source["path"],
                    "sheet_name": sheet.title,
                    "source_row_number": row_index,
                    "visit_date": visit_date_text,
                    "driver_name": clean_text(padded[0]),
                    "planned_time": time_text(padded[1]),
                    "professionals": professionals,
                    "patient_name": patient_name,
                    "service_text": service_text,
                    "address_text": address_text,
                    "phone_text": clean_text(padded[10]),
                    "raw_record": {"row": [clean_text(value) for value in padded]},
                }
            )
    return rows


def workbook_metadata(path: Path) -> dict[str, Any]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    return {
        "sheet_count": len(workbook.worksheets),
        "sheets": [{"title": sheet.title, "max_row": sheet.max_row, "max_column": sheet.max_column} for sheet in workbook.worksheets],
    }


def drive_file_record(source: dict[str, Any], folder_url: str, dataset: str, local_path: Path | None = None, metadata: dict[str, Any] | None = None) -> dict[str, Any]:
    size = local_path.stat().st_size if local_path and local_path.exists() else None
    checksum = sha256(local_path) if local_path and local_path.exists() else None
    month = source.get("month")
    year = source.get("year")
    return {
        "drive_id": source["drive_id"],
        "folder_url": folder_url,
        "path": source["path"],
        "title": source.get("title") or Path(source["path"]).name,
        "mime_type": source.get("mime_type", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        "dataset": dataset,
        "year": year,
        "month": month,
        "period": f"{year}-{month:02d}" if year and month else None,
        "local_path": str(local_path) if local_path else None,
        "sha256": checksum,
        "size_bytes": size,
        "metadata": metadata or {},
    }


def metrics_folder_files() -> list[dict[str, Any]]:
    consolidation = json.loads(CONSOLIDATION_SPEC.read_text(encoding="utf-8"))
    by_id = {item["id"]: item for item in METRICS_MANIFEST}
    files = []
    for item in consolidation["files"]:
        source = by_id.get(item["id"], item)
        source = {"drive_id": item["id"], **source, **item}
        files.append(
            {
                "drive_id": item["id"],
                "folder_url": METRICS_FOLDER_URL,
                "path": item["path"],
                "title": item["title"],
                "mime_type": item["mime_type"],
                "dataset": item["dataset"],
                "year": item.get("year"),
                "month": item.get("month"),
                "period": f"{item['year']}-{item['month']:02d}" if item.get("year") and item.get("month") else None,
                "local_path": item.get("local_file"),
                "sha256": item.get("sha256"),
                "size_bytes": item.get("size_bytes"),
                "metadata": {
                    "workbook": item.get("workbook"),
                    "source_url": source_url(item["id"], source.get("native_google", False)),
                },
            }
        )
    return files


def build_annual_metric_rows() -> list[dict[str, Any]]:
    payload = json.loads(METRICS_SPEC.read_text(encoding="utf-8"))
    quality_flags = payload.get("quality_flags", [])
    rows: list[dict[str, Any]] = []
    unique_by_year = {row["year"]: row for row in payload.get("unique_user_rows", [])}
    for year_text, metrics in payload.get("annual_rem", {}).items():
        year = int(year_text)
        for name in ("atenciones_rem_visitas", "usuarios_rem_personas_atendidas_suma_mensual", "ingresos_rem", "dias_persona"):
            rows.append(
                {
                    "metric_id": make_id("drv_metric", f"{year}|{name}|rem"),
                    "year": year,
                    "metric_name": name,
                    "metric_value": metrics.get(name),
                    "source_type": "rem",
                    "source_detail": "REM A21 C.1 Drive",
                    "quality_flags": quality_flags,
                    "metadata": {"periods": metrics.get("periods"), "visitas_by_profession": metrics.get("visitas_by_profession")},
                }
            )
        unique = unique_by_year.get(year)
        if unique:
            rows.append(
                {
                    "metric_id": make_id("drv_metric", f"{year}|usuarios_unicos|{unique.get('source')}"),
                    "year": year,
                    "metric_name": "usuarios_unicos_disponibles",
                    "metric_value": unique.get("unique_users"),
                    "source_type": "nominal_or_db",
                    "source_detail": unique.get("source"),
                    "quality_flags": quality_flags,
                    "metadata": {"note": unique.get("note")},
                }
            )
    return rows


def build_rem_snapshots() -> list[dict[str, Any]]:
    consolidation = json.loads(CONSOLIDATION_SPEC.read_text(encoding="utf-8"))
    rem_file_by_period = {
        f"{item.get('year')}-{item.get('month'):02d}": item
        for item in consolidation["files"]
        if item.get("dataset") == "rem" and item.get("year") and item.get("month")
    }
    rows = []
    for metrics in consolidation["metrics"].get("rem_metrics", []):
        period = metrics["period"]
        source_file = rem_file_by_period.get(period, {})
        source_id = f"drive_{source_file.get('id', period)}"
        rows.append(
            {
                "snapshot_id": make_id("rem_snapshot", f"drive|{period}"),
                "source_id": source_id,
                "source_file": source_file,
                "periodo": period,
                "payload": metrics,
                "criterios": {"source": "Drive REM A21 C.1", "extractor": "scripts/drive_consolidation.py"},
                "calidad_datos": {
                    "flags": [flag for flag in consolidation["metrics"].get("quality_flags", []) if flag.get("period") == period]
                },
            }
        )
    return rows


def parse_handover_dates(title: str) -> tuple[str | None, str | None]:
    normalized = normalize_text(title)
    years = [int(value) for value in re.findall(r"\b(20\d{2})\b", normalized)]
    if not years:
        return None, None
    start_year = years[0]
    end_year = years[-1]
    month_matches = [(match.start(), MONTHS[match.group(1)]) for match in re.finditer(r"\b(" + "|".join(MONTHS) + r")\b", normalized)]
    numbers = [(match.start(), int(match.group(1))) for match in re.finditer(r"\b([0-3]?\d)\b", normalized) if 1 <= int(match.group(1)) <= 31]
    if len(numbers) < 2 or not month_matches:
        return None, None
    start_day = numbers[0][1]
    end_day = numbers[1][1]
    start_month = month_matches[0][1]
    end_month = month_matches[-1][1] if len(month_matches) > 1 else start_month
    try:
        start = date(start_year, start_month, start_day).isoformat()
        end = date(end_year, end_month, end_day).isoformat()
    except ValueError:
        return None, None
    return start, end


def handover_local_path(directory: Path, item: dict[str, Any]) -> Path:
    raw_name = Path(item["path"]).name
    if not raw_name.lower().endswith(".docx"):
        raw_name += ".docx"
    return directory / raw_name


def download_handover_file(item: dict[str, Any], directory: Path, force: bool = False) -> Path:
    directory.mkdir(parents=True, exist_ok=True)
    target = handover_local_path(directory, item)
    if target.exists() and target.stat().st_size > 0 and not force:
        return target
    subprocess.run(
        ["curl", "-L", "-s", "-o", str(target), f"https://drive.google.com/uc?export=download&id={item['id']}"],
        check=True,
    )
    return target


def load_handover_manifest(manifest_path: Path, directory: Path, force_download: bool = False) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    if not manifest_path.exists():
        return load_handover_downloads(directory)
    items = json.loads(manifest_path.read_text(encoding="utf-8"))
    files: list[dict[str, Any]] = []
    handovers: list[dict[str, Any]] = []
    for item in items:
        path = download_handover_file(item, directory, force=force_download)
        source_file, handover = parse_handover_file(path, item["id"], item["path"])
        files.append(source_file)
        handovers.append(handover)
    return files, handovers


def parse_handover_file(path: Path, drive_id: str, source_path: str) -> tuple[dict[str, Any], dict[str, Any]]:
    title = Path(source_path).stem
    checksum = sha256(path)
    period_start, period_end = parse_handover_dates(title)
    source = {
        "drive_id": drive_id,
        "path": source_path if source_path.lower().endswith(".docx") else f"{source_path}.docx",
        "title": title,
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    file_record = drive_file_record(
        source,
        HANDOVER_FOLDER_URL,
        "entrega_turno",
        path,
        {"source_note": "Downloaded from public Drive folder", "source_url": source_url(drive_id)},
    )
    try:
        document = Document(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    except Exception as exc:  # pragma: no cover - defensive for malformed docs
        text = ""
        metadata = {"parse_error": str(exc), "sha256": checksum}
    else:
        metadata = {"paragraph_count": len(document.paragraphs), "sha256": checksum}
    handover = {
        "handover_id": make_id("handover", f"{drive_id}|{title}"),
        "drive_id": drive_id,
        "source_path": source["path"],
        "title": title,
        "period_start": period_start,
        "period_end": period_end,
        "text_content": text,
        "metadata": metadata,
    }
    return file_record, handover


def load_handover_downloads(directory: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    files: list[dict[str, Any]] = []
    handovers: list[dict[str, Any]] = []
    if not directory.exists():
        return files, handovers
    for path in sorted(directory.glob("*.docx")):
        drive_id = make_id("handover_file", f"{path.stem}|{sha256(path)}")
        source_file, handover = parse_handover_file(path, drive_id, path.name)
        files.append(source_file)
        handovers.append(handover)
    return files, handovers


def insert_drive_source_file_sql(row: dict[str, Any]) -> str:
    cols = ["drive_id", "folder_url", "path", "title", "mime_type", "dataset", "year", "month", "period", "local_path", "sha256", "size_bytes", "metadata"]
    values = [json_literal(row[column]) if column == "metadata" else sql_literal(row.get(column)) for column in cols]
    return (
        f"INSERT INTO staging.drive_source_file ({', '.join(cols)}) VALUES ({', '.join(values)}) "
        "ON CONFLICT (drive_id) DO UPDATE SET "
        "folder_url=EXCLUDED.folder_url, path=EXCLUDED.path, title=EXCLUDED.title, mime_type=EXCLUDED.mime_type, "
        "dataset=EXCLUDED.dataset, year=EXCLUDED.year, month=EXCLUDED.month, period=EXCLUDED.period, "
        "local_path=EXCLUDED.local_path, sha256=EXCLUDED.sha256, size_bytes=EXCLUDED.size_bytes, "
        "metadata=EXCLUDED.metadata, updated_at=now();"
    )


def insert_original_source_sql(row: dict[str, Any]) -> str:
    source_id = f"drive_{row['drive_id']}"
    url = row.get("metadata", {}).get("source_url") or source_url(row["drive_id"], row["mime_type"] == "application/vnd.google-apps.spreadsheet")
    return (
        "INSERT INTO reference.original_source "
        "(source_id, title, authority, source_type, version_label, source_url, local_path, checksum_sha256, mime_type, notes) VALUES "
        f"({sql_literal(source_id)}, {sql_literal(row['title'])}, 'Hospital de San Carlos Dr. Benicio Arzola Medina - HODOM', "
        f"'drive_source', {sql_literal(row.get('period'))}, {sql_literal(url)}, {sql_literal(row.get('local_path'))}, "
        f"{sql_literal(row.get('sha256'))}, {sql_literal(row.get('mime_type'))}, {sql_literal('Registered by drive staging migration 2026-05-25')}) "
        "ON CONFLICT (source_id) DO UPDATE SET title=EXCLUDED.title, source_url=EXCLUDED.source_url, "
        "local_path=EXCLUDED.local_path, checksum_sha256=EXCLUDED.checksum_sha256, mime_type=EXCLUDED.mime_type, retrieved_at=now();"
    )


def insert_route_visit_sql(row: dict[str, Any]) -> str:
    cols = [
        "route_visit_id",
        "drive_id",
        "source_path",
        "sheet_name",
        "source_row_number",
        "visit_date",
        "driver_name",
        "planned_time",
        "professionals",
        "patient_name",
        "service_text",
        "address_text",
        "phone_text",
        "raw_record",
    ]
    values = [json_literal(row[column]) if column in {"professionals", "raw_record"} else sql_literal(row.get(column)) for column in cols]
    return (
        f"INSERT INTO staging.hodom_route_visit ({', '.join(cols)}) VALUES ({', '.join(values)}) "
        "ON CONFLICT (drive_id, sheet_name, source_row_number) DO UPDATE SET "
        "visit_date=EXCLUDED.visit_date, driver_name=EXCLUDED.driver_name, planned_time=EXCLUDED.planned_time, "
        "professionals=EXCLUDED.professionals, patient_name=EXCLUDED.patient_name, service_text=EXCLUDED.service_text, "
        "address_text=EXCLUDED.address_text, phone_text=EXCLUDED.phone_text, raw_record=EXCLUDED.raw_record, updated_at=now();"
    )


def insert_handover_sql(row: dict[str, Any]) -> str:
    cols = ["handover_id", "drive_id", "source_path", "title", "period_start", "period_end", "text_content", "metadata"]
    values = [json_literal(row[column]) if column == "metadata" else sql_literal(row.get(column)) for column in cols]
    return (
        f"INSERT INTO staging.hodom_shift_handover ({', '.join(cols)}) VALUES ({', '.join(values)}) "
        "ON CONFLICT (handover_id) DO UPDATE SET "
        "period_start=EXCLUDED.period_start, period_end=EXCLUDED.period_end, text_content=EXCLUDED.text_content, "
        "metadata=EXCLUDED.metadata, updated_at=now();"
    )


def insert_annual_metric_sql(row: dict[str, Any]) -> str:
    cols = ["metric_id", "year", "metric_name", "metric_value", "source_type", "source_detail", "quality_flags", "metadata"]
    values = [json_literal(row[column]) if column in {"quality_flags", "metadata"} else sql_literal(row.get(column)) for column in cols]
    return (
        f"INSERT INTO staging.hodom_annual_metric ({', '.join(cols)}) VALUES ({', '.join(values)}) "
        "ON CONFLICT (year, metric_name, source_type) DO UPDATE SET "
        "metric_value=EXCLUDED.metric_value, source_detail=EXCLUDED.source_detail, quality_flags=EXCLUDED.quality_flags, "
        "metadata=EXCLUDED.metadata, updated_at=now();"
    )


def insert_rem_snapshot_sql(row: dict[str, Any]) -> str:
    return (
        "INSERT INTO reporting.rem_a21_c1_snapshot "
        "(snapshot_id, source_id, periodo, establecimiento_nombre, payload, criterios, calidad_datos) VALUES "
        f"({sql_literal(row['snapshot_id'])}, {sql_literal(row['source_id'])}, {sql_literal(row['periodo'])}, "
        "'Hospital de San Carlos Dr. Benicio Arzola Medina', "
        f"{json_literal(row['payload'])}, {json_literal(row['criterios'])}, {json_literal(row['calidad_datos'])}) "
        "ON CONFLICT (snapshot_id) DO UPDATE SET payload=EXCLUDED.payload, criterios=EXCLUDED.criterios, "
        "calidad_datos=EXCLUDED.calidad_datos, updated_at=now();"
    )


def provenance_sql(target_table: str, target_pk: str, source_file: str, source_key: str | None = None) -> str:
    return (
        "INSERT INTO migration.provenance (target_table, target_pk, source_type, source_file, source_key, phase, field_name) VALUES "
        f"({sql_literal(target_table)}, {sql_literal(target_pk)}, 'google_drive', {sql_literal(source_file)}, "
        f"{sql_literal(source_key)}, {sql_literal(PHASE)}, NULL);"
    )


def build_sql_batch(
    files: list[dict[str, Any]],
    route_rows: list[dict[str, Any]],
    handovers: list[dict[str, Any]],
    annual_metrics: list[dict[str, Any]],
    rem_snapshots: list[dict[str, Any]],
) -> str:
    lines = [
        "-- Generated by scripts/drive_staging_migration.py",
        "BEGIN;",
        f"DELETE FROM migration.provenance WHERE phase = {sql_literal(PHASE)};",
    ]
    for row in files:
        lines.append(insert_original_source_sql(row))
    for row in files:
        lines.append(insert_drive_source_file_sql(row))
        lines.append(provenance_sql("staging.drive_source_file", row["drive_id"], row["path"], row.get("sha256")))
    for row in route_rows:
        lines.append(insert_route_visit_sql(row))
        lines.append(provenance_sql("staging.hodom_route_visit", row["route_visit_id"], row["source_path"], f"{row['sheet_name']}:{row['source_row_number']}"))
    for row in handovers:
        lines.append(insert_handover_sql(row))
        lines.append(provenance_sql("staging.hodom_shift_handover", row["handover_id"], row["source_path"], row.get("period_start")))
    for row in annual_metrics:
        lines.append(insert_annual_metric_sql(row))
        lines.append(provenance_sql("staging.hodom_annual_metric", row["metric_id"], METRICS_SPEC.name, str(row["year"])))
    for row in rem_snapshots:
        lines.append(insert_rem_snapshot_sql(row))
        lines.append(provenance_sql("reporting.rem_a21_c1_snapshot", row["snapshot_id"], row.get("source_file", {}).get("path") or "REM Drive", row["periodo"]))
    lines.append("COMMIT;")
    return "\n".join(lines) + "\n"


def build_payload(force_download: bool = False, handover_dir: Path | None = None, handover_manifest: Path | None = None) -> dict[str, Any]:
    files = metrics_folder_files()
    route_rows: list[dict[str, Any]] = []
    for source in ROUTE_SOURCES:
        local_path = download_drive_file(source, force=force_download)
        metadata = workbook_metadata(local_path)
        file_record = drive_file_record(
            {
                **source,
                "mime_type": "application/vnd.google-apps.spreadsheet" if source.get("native_google") else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            },
            ROUTES_FOLDER_URL,
            "rutas",
            local_path,
            metadata,
        )
        files.append(file_record)
        route_rows.extend(parse_route_workbook(local_path, source))

    handover_files, handovers = load_handover_manifest(
        handover_manifest or HANDOVER_MANIFEST,
        handover_dir or OUT_DIR / "handover-download",
        force_download=force_download,
    )
    files.extend(handover_files)
    annual_metrics = build_annual_metric_rows()
    rem_snapshots = build_rem_snapshots()
    return {
        "files": files,
        "route_rows": route_rows,
        "handovers": handovers,
        "annual_metrics": annual_metrics,
        "rem_snapshots": rem_snapshots,
    }


def apply_sql(db_url: str, sql_path: Path) -> None:
    subprocess.run(["psql", db_url, "-v", "ON_ERROR_STOP=1", "-f", str(sql_path)], check=True)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--apply", action="store_true", help="Apply generated load SQL to PostgreSQL.")
    parser.add_argument("--force-download", action="store_true")
    parser.add_argument("--db-url", default=DB_URL)
    parser.add_argument("--handover-dir", type=Path, default=OUT_DIR / "handover-download")
    parser.add_argument("--handover-manifest", type=Path, default=HANDOVER_MANIFEST)
    args = parser.parse_args()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    payload = build_payload(force_download=args.force_download, handover_dir=args.handover_dir, handover_manifest=args.handover_manifest)
    sql = build_sql_batch(**payload)
    LOAD_SQL.write_text(sql, encoding="utf-8")
    summary = {key: len(value) for key, value in payload.items()}
    summary["load_sql"] = str(LOAD_SQL)
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    if args.apply:
        apply_sql(args.db_url, LOAD_SQL)


if __name__ == "__main__":
    main()
